"""Convert 寻道 分镜头脚本 Markdown to formatted DOCX."""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_from_markdown(doc, lines, start_idx):
    """Parse a markdown table and add it to the document."""
    # Find table lines
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1

    if len(table_lines) < 2:
        return i

    # Parse rows
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)

    # Filter separator row
    data_rows = [r for r in rows if not all(re.match(r'^[-:]+$', c) for c in r)]

    if not data_rows:
        return i

    ncols = len(data_rows[0])
    nrows = len(data_rows)

    table = doc.add_table(rows=nrows, cols=ncols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < ncols:
                cell = table.rows[r_idx].cells[c_idx]
                # Clear default paragraph
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(cell_text)
                run.font.size = Pt(8)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                if r_idx == 0:
                    set_cell_shading(cell, '2d2d2d')
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True

    doc.add_paragraph()  # spacing after table
    return i


def process_markdown_to_docx(md_path, docx_path):
    """Main conversion function."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    i = 0
    in_code_block = False
    in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Table detection
        if stripped.startswith('|') and not in_code_block:
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|') and re.match(r'^[\|\s\-:]+$', lines[i + 1].strip().replace('|', '').replace('-', '').replace(':', '').replace(' ', '')):
                i = add_table_from_markdown(doc, lines, i)
                continue
            i += 1
            continue

        # Title page (# 寻道)
        if stripped.startswith('# ') and '寻道' in stripped:
            title_text = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(72)
            run = p.add_run(title_text)
            run.font.size = Pt(26)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.color.rgb = RGBColor(0x8b, 0x00, 0x00)

            # Subtitle
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run('分镜头脚本 · 第二至五集')
            run2.font.size = Pt(14)
            run2.font.name = 'Microsoft YaHei'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            doc.add_page_break()
            i += 1
            continue

        # Episode headers (e.g., # 《寻道》第二集《运势》)
        if stripped.startswith('# ') and '第' in stripped and '集' in stripped:
            doc.add_page_break()
            p = doc.add_paragraph()
            p.space_before = Pt(6)
            run = p.add_run(stripped[2:])
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.color.rgb = RGBColor(0x8b, 0x00, 0x00)
            i += 1
            continue

        # Section headers (##)
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            p.space_before = Pt(18)
            p.space_after = Pt(6)
            run = p.add_run(stripped[3:])
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
            i += 1
            continue

        # Shot headers (###)
        if stripped.startswith('### '):
            p = doc.add_paragraph()
            p.space_before = Pt(14)
            p.space_after = Pt(4)
            run = p.add_run(stripped[4:])
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.space_before = Pt(4)
            p.space_after = Pt(4)
            run = p.add_run('─' * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
            i += 1
            continue

        # Code blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(line)
            run.font.size = Pt(8)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Blockquotes
        if stripped.startswith('> '):
            quote_text = stripped[2:]
            # Remove bold markers for processing
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.2)
            p.space_before = Pt(2)
            p.space_after = Pt(2)

            # Parse inline bold
            parts = re.split(r'(\*\*.*?\*\*)', quote_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(9)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1

            # Handle multi-line blockquotes without > prefix
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('>') and not lines[i].strip().startswith('#') and lines[i].strip() != '---' and not lines[i].strip().startswith('|'):
                cont = lines[i]
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Cm(1.2)
                run2 = p2.add_run(cont.strip())
                run2.font.size = Pt(9)
                run2.font.name = 'Microsoft YaHei'
                run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                i += 1
            continue

        # Bold list items (e.g., - **term**: description)
        if stripped.startswith('- **') or stripped.startswith('  - **'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.space_before = Pt(1)
            p.space_after = Pt(1)

            # Parse the bold part
            match = re.match(r'(\s*-\s+)\*\*(.+?)\*\*(.*)', stripped)
            if match:
                prefix = match.group(1)
                bold_text = match.group(2)
                rest = match.group(3)

                run_prefix = p.add_run(prefix)
                run_prefix.font.size = Pt(9.5)
                run_prefix.font.name = 'Microsoft YaHei'
                run_prefix._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

                run_bold = p.add_run(bold_text)
                run_bold.font.bold = True
                run_bold.font.size = Pt(9.5)
                run_bold.font.name = 'Microsoft YaHei'
                run_bold._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

                if rest:
                    run_rest = p.add_run(rest)
                    run_rest.font.size = Pt(9.5)
                    run_rest.font.name = 'Microsoft YaHei'
                    run_rest._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            else:
                run = p.add_run(stripped)
                run.font.size = Pt(9.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            i += 1
            continue

        # Regular list items
        if stripped.startswith('- ') or stripped.startswith('  - '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.space_before = Pt(1)
            p.space_after = Pt(1)

            # Handle inline bold
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(9.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            i += 1
            continue

        # Image lines
        if stripped.startswith('16:9'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(stripped)
            run.font.size = Pt(7)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.space_before = Pt(1)
        p.space_after = Pt(1)

        # Handle inline bold markers
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        i += 1

    doc.save(docx_path)
    print(f"DOCX saved to: {docx_path}")

if __name__ == '__main__':
    md_path = r'C:\Users\Brian\Desktop\寻道_第二至五集_分镜头脚本.md'
    docx_path = r'C:\Users\Brian\Desktop\寻道_第二至五集_分镜头脚本.docx'
    process_markdown_to_docx(md_path, docx_path)
